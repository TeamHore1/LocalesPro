#!/usr/bin/env python3
"""Convert Locales Pro BlackBox Markdown files to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BLACKBOX = ROOT / "BlackBox"

TITLE_COLOR = RGBColor(9, 35, 121)
ACCENT_COLOR = RGBColor(37, 99, 235)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_styled_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(text, level=min(level, 3))
    for run in heading.runs:
        run.font.color.rgb = TITLE_COLOR if level == 1 else ACCENT_COLOR


def parse_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        content = match.group(1) or match.group(2)
        run = paragraph.add_run(content)
        run.bold = bool(match.group(1))
        if match.group(2):
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph(doc: Document, text: str) -> None:
    text = text.strip()
    if not text:
        return
    paragraph = doc.add_paragraph()
    parse_inline_runs(paragraph, text)


def add_bullet(doc: Document, text: str, ordered: bool = False) -> None:
    paragraph = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    parse_inline_runs(paragraph, text.strip().lstrip("- ").lstrip("* "))


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    if row_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if row_idx == 0:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): "092379",
                    },
                )
                shading.append(shd)

    doc.add_paragraph()


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_defaults(doc)

    title = md_path.stem.replace("-", " ")
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("Locales Pro")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(title)
    sub_run.bold = True
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = TITLE_COLOR

    doc.add_paragraph()

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            index += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            add_styled_heading(doc, text, level)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            table_rows = [parse_table_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[index]))
                index += 1
            add_table(doc, table_rows)
            continue

        if stripped.startswith("|"):
            table_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                if not is_table_separator(lines[index]):
                    table_rows.append(parse_table_row(lines[index]))
                index += 1
            add_table(doc, table_rows)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet(doc, stripped)
            index += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_bullet(doc, re.sub(r"^\d+\.\s", "", stripped), ordered=True)
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        add_paragraph(doc, stripped)
        index += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Created: {docx_path}")


def build_combined_docx(output_path: Path) -> None:
    order = [
        "README.md",
        "01-Blackbox-Pembayaran.md",
        "02-Boundary-Value-Analysis.md",
        "03-Equivalence-Partitioning.md",
        "04-Use-Case-Testing.md",
        "05-Use-Case-Diagram-Transaksi.md",
    ]

    doc = Document()
    set_document_defaults(doc)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("Locales Pro")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT_COLOR

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Dokumen BlackBox Testing Lengkap")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = TITLE_COLOR

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Versi 2.0 | Pembayaran Tunai/Manual | Tanpa Midtrans/QRIS").font.size = Pt(11)

    doc.add_page_break()

    for filename in order:
        md_path = BLACKBOX / filename
        if not md_path.exists():
            continue

        temp_docx = BLACKBOX / f"_temp_{filename.replace('.md', '.docx')}"
        markdown_to_docx(md_path, temp_docx)

        part = Document(temp_docx)
        for element in part.element.body:
            doc.element.body.append(element)
        doc.add_page_break()
        temp_docx.unlink(missing_ok=True)

    if doc.paragraphs and doc.paragraphs[-1].text == "":
        pass

    doc.save(output_path)
    print(f"Created: {output_path}")


def main() -> int:
    files = [
        ("README.md", "00-Indeks-BlackBox.docx"),
        ("01-Blackbox-Pembayaran.md", "01-Blackbox-Pembayaran.docx"),
        ("02-Boundary-Value-Analysis.md", "02-Boundary-Value-Analysis.docx"),
        ("03-Equivalence-Partitioning.md", "03-Equivalence-Partitioning.docx"),
        ("04-Use-Case-Testing.md", "04-Use-Case-Testing.docx"),
        ("05-Use-Case-Diagram-Transaksi.md", "05-Use-Case-Diagram-Transaksi.docx"),
    ]

    for md_name, docx_name in files:
        md_path = BLACKBOX / md_name
        if md_path.exists():
            markdown_to_docx(md_path, BLACKBOX / docx_name)

    build_combined_docx(BLACKBOX / "Blackbox-Pembayaran-Lengkap.docx")
    return 0


if __name__ == "__main__":
    sys.exit(main())
