from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "SKPL-assets"
REPORT_MD = ROOT / "SKPL- LocalPro - lengkap.md"
REPORT_DOCX = ROOT / "SKPL- LocalPro - lengkap.docx"
REPORT_DATE = "25 Mei 2026"


def load_font(size=26, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
        if bold
        else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


FONT_TITLE = load_font(42, True)
FONT_SUBTITLE = load_font(24, False)
FONT_H = load_font(24, True)
FONT = load_font(20, False)
FONT_SMALL = load_font(17, False)
FONT_TINY = load_font(15, False)


def text_width(draw, text, font):
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0]


def wrap_text(draw, text, font, max_width):
    words = str(text).split()
    if not words:
        return [""]
    lines = []
    current = words[0]
    for word in words[1:]:
        test = f"{current} {word}"
        if text_width(draw, test, font) <= max_width:
            current = test
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_box(draw, xy, title, lines, fill, outline, title_fill=None, font=FONT_SMALL):
    x1, y1, x2, y2 = xy
    shadow = 8
    draw.rectangle((x1 + shadow, y1 + shadow, x2 + shadow, y2 + shadow), fill="#dfe5ef")
    draw.rectangle(xy, fill=fill, outline=outline, width=3)
    header_h = 42
    draw.rectangle((x1, y1, x2, y1 + header_h), fill=title_fill or "#e9eef8", outline=outline, width=3)
    draw.text((x1 + 14, y1 + 10), title, fill=outline, font=FONT_H)
    current_y = y1 + header_h + 12
    for line in lines:
        for wrapped in wrap_text(draw, line, font, x2 - x1 - 28):
            draw.text((x1 + 14, current_y), wrapped, fill="#1f2937", font=font)
            current_y += 24


def draw_ellipse(draw, xy, label, fill="#f8fafc", outline="#092379"):
    draw.ellipse(xy, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = xy
    lines = wrap_text(draw, label, FONT_SMALL, x2 - x1 - 30)
    total_h = len(lines) * 24
    current_y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = text_width(draw, line, FONT_SMALL)
        draw.text((x1 + ((x2 - x1) - w) / 2, current_y), line, fill=outline, font=FONT_SMALL)
        current_y += 24


def draw_arrow(draw, start, end, fill="#092379", width=4, label=None, font=FONT_TINY):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    head_len = 16
    head_angle = math.pi / 7
    points = [
        (x2, y2),
        (x2 - head_len * math.cos(angle - head_angle), y2 - head_len * math.sin(angle - head_angle)),
        (x2 - head_len * math.cos(angle + head_angle), y2 - head_len * math.sin(angle + head_angle)),
    ]
    draw.polygon(points, fill=fill)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        draw.rectangle((mx - 8, my - 18, mx + text_width(draw, label, font) + 8, my + 12), fill="#ffffff")
        draw.text((mx, my - 14), label, fill=fill, font=font)


def base_canvas(title, subtitle, size=(1800, 1080)):
    img = Image.new("RGB", size, "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, size[0], 70), fill="#dbe7f7")
    draw.text((34, 16), title, fill="#092379", font=FONT_TITLE)
    draw.text((34, 88), subtitle, fill="#64748b", font=FONT_SUBTITLE)
    draw.rectangle((20, 20, size[0] - 20, size[1] - 20), outline="#4b5563", width=3)
    return img, draw


def generate_use_case_diagram():
    img, draw = base_canvas(
        "DIAGRAM USE CASE LOCALESPRO",
        "Aktor dan fitur sesuai aplikasi POS tunai/manual saat ini",
    )
    draw_box(draw, (70, 210, 300, 335), "AKTOR", ["Admin"], "#f1f5f9", "#092379")
    draw_box(draw, (70, 475, 300, 600), "AKTOR", ["Kasir"], "#fff7ed", "#b45309")
    draw_box(draw, (70, 740, 300, 865), "AKTOR", ["Pembeli"], "#ecfdf5", "#047857")

    draw.rectangle((390, 155, 1450, 1000), fill="#ffffff", outline="#092379", width=4)
    draw.text((780, 175), "SISTEM LOCALESPRO", fill="#092379", font=FONT_H)

    use_cases = [
        ((455, 245, 700, 340), "Login\nA/K"),
        ((760, 245, 1005, 340), "Registrasi Kasir\nK"),
        ((1065, 245, 1310, 340), "Review Akun Kasir\nA"),
        ((455, 405, 700, 500), "Dashboard\nA"),
        ((760, 405, 1005, 500), "Kelola Cabang\nA"),
        ((1065, 405, 1310, 500), "Kelola Produk & Resep\nA"),
        ((455, 565, 700, 660), "Kelola Bahan\nA"),
        ((760, 565, 1005, 660), "Update Stok\nA"),
        ((1065, 565, 1310, 660), "Riwayat Mutasi Stok\nA/K"),
        ((455, 730, 700, 825), "POS Tunai\nA/K/P"),
        ((760, 730, 1005, 825), "Cetak Struk\nK/P"),
        ((1065, 730, 1310, 825), "Laporan & Void\nA/K"),
    ]
    for xy, label in use_cases:
        draw_ellipse(draw, xy, label)

    draw_arrow(draw, (300, 272), (390, 272), "#092379", 4, label="akses admin")
    draw_arrow(draw, (300, 535), (390, 535), "#b45309", 4, label="akses kasir")
    draw_arrow(draw, (300, 802), (390, 802), "#047857", 4, label="transaksi")
    draw_box(
        draw,
        (1490, 245, 1730, 470),
        "LEGENDA",
        ["A = Admin", "K = Kasir", "P = Pembeli"],
        "#ffffff",
        "#64748b",
        font=FONT_TINY,
    )
    draw_box(
        draw,
        (1490, 565, 1730, 785),
        "CATATAN",
        ["Hak akses detail juga tersedia pada tabel aktor dan tabel kebutuhan fungsional."],
        "#ffffff",
        "#64748b",
        font=FONT_TINY,
    )

    img.save(ASSET_DIR / "diagram-use-case.png")


def generate_architecture_diagram():
    img, draw = base_canvas(
        "DIAGRAM ARSITEKTUR LOCALESPRO",
        "React/Vite SPA, PHP API, MySQL/MariaDB, dan penyimpanan sesi browser",
    )
    draw_box(
        draw,
        (80, 230, 340, 430),
        "PENGGUNA",
        ["Admin", "Kasir", "Pembeli"],
        "#fff7ed",
        "#b45309",
    )
    draw_box(
        draw,
        (470, 160, 820, 470),
        "FRONTEND",
        ["React Vite SPA", "Login dan route guard", "Dashboard", "POS", "Master data", "Laporan"],
        "#eff6ff",
        "#092379",
    )
    draw_box(
        draw,
        (470, 625, 820, 820),
        "LOCAL STORAGE",
        ["JWT token", "Data user", "Cabang aktif", "Masa sesi"],
        "#f8fafc",
        "#64748b",
    )
    draw_box(
        draw,
        (1000, 160, 1350, 560),
        "BACKEND API",
        ["PHP endpoint per modul", "Auth JWT dan role", "Branches users", "Products ingredients", "Transactions", "Stock movements"],
        "#ecfdf5",
        "#047857",
    )
    draw_box(
        draw,
        (1460, 160, 1730, 420),
        "DATABASE",
        ["MySQL/MariaDB", "branches users", "products recipes", "transactions items", "stock_movements"],
        "#f5f3ff",
        "#6d28d9",
    )
    draw_box(
        draw,
        (1000, 680, 1350, 870),
        "FILE & LOG",
        ["Gambar produk base64", "auth_security.log", "login_attempts.json", "config lokal"],
        "#fef2f2",
        "#b91c1c",
    )
    draw_arrow(draw, (340, 320), (470, 320), "#b45309", label="HTTP browser")
    draw_arrow(draw, (820, 320), (1000, 320), "#092379", label="JSON API")
    draw_arrow(draw, (1350, 285), (1460, 285), "#047857", label="PDO SQL")
    draw_arrow(draw, (645, 470), (645, 625), "#64748b", label="simpan sesi")
    draw_arrow(draw, (1175, 560), (1175, 680), "#b91c1c", label="audit")
    draw_box(
        draw,
        (960, 900, 1730, 1010),
        "CATATAN",
        ["Aplikasi saat ini tidak memakai payment gateway eksternal; transaksi diproses sebagai Cash/manual."],
        "#ffffff",
        "#64748b",
        font=FONT_TINY,
    )
    img.save(ASSET_DIR / "diagram-architecture.png")


def generate_sequence_diagram():
    img, draw = base_canvas(
        "SEQUENCE DIAGRAM POS TUNAI",
        "Alur transaksi tunai dari pilih item sampai cetak struk",
        size=(1900, 1120),
    )
    actors = [
        ("Kasir", 170, "#b45309"),
        ("Frontend POS", 550, "#092379"),
        ("Backend API", 970, "#047857"),
        ("MySQL", 1360, "#6d28d9"),
        ("Printer/Browser", 1710, "#64748b"),
    ]
    y_top, y_bottom = 170, 995
    for name, x, color in actors:
        draw_box(draw, (x - 120, 125, x + 120, 205), name.upper(), [], "#ffffff", color)
        draw.line((x, 205, x, y_bottom), fill=color, width=2)

    steps = [
        (170, 550, 260, "1. Pilih produk dan qty", "#b45309"),
        (550, 550, 330, "2. Hitung total dan cek stok resep", "#092379"),
        (550, 970, 410, "3. POST /transactions/create.php", "#092379"),
        (970, 970, 485, "4. Validasi token, role, cabang, item", "#047857"),
        (970, 1360, 565, "5. Insert transactions", "#047857"),
        (970, 1360, 635, "6. Insert transaction_items", "#047857"),
        (970, 1360, 705, "7. Deduct ingredients dan record stock_movements", "#047857"),
        (1360, 970, 775, "8. Commit success", "#6d28d9"),
        (970, 550, 850, "9. Return id dan transaction_code", "#047857"),
        (550, 1710, 925, "10. Buka jendela struk dan print", "#64748b"),
    ]
    for x1, x2, y, label, color in steps:
        if x1 == x2:
            draw.arc((x1 - 55, y - 28, x1 + 55, y + 28), start=90, end=270, fill=color, width=4)
            draw.text((x1 + 65, y - 12), label, fill=color, font=FONT_TINY)
        else:
            draw_arrow(draw, (x1, y), (x2, y), color, 4, label=label, font=FONT_TINY)
    draw_box(
        draw,
        (460, 1010, 1440, 1090),
        "CATATAN",
        ["Pembayaran tunai langsung berstatus Paid; stok bahan otomatis berkurang berdasarkan resep dan dikembalikan jika transaksi di-void."],
        "#f8fafc",
        "#64748b",
        font=FONT_TINY,
    )
    img.save(ASSET_DIR / "diagram-sequence-pos-cash.png")


def generate_data_model_diagram():
    img, draw = base_canvas(
        "DATA MODEL LOCALESPRO",
        "Entitas utama dan relasi sesuai skema database aplikasi saat ini",
        size=(2100, 1350),
    )
    boxes = {
        "branches": (80, 175, 485, 395, "#eff6ff", "#092379", ["PK id", "name, address, phone", "status, created_at"]),
        "users": (590, 175, 995, 455, "#ecfdf5", "#047857", ["PK id", "username, password", "full_name, email, phone", "role, status", "FK branch_id", "approved_by, approved_at"]),
        "transactions": (1100, 175, 1530, 455, "#f5f3ff", "#6d28d9", ["PK id", "transaction_code", "FK user_id, branch_id", "total_price", "payment_method", "payment_status", "paid_at"]),
        "transaction_items": (1615, 175, 2045, 455, "#fef2f2", "#b91c1c", ["PK id", "FK transaction_id", "FK product_id", "quantity", "subtotal"]),
        "ingredients": (80, 650, 520, 930, "#fff7ed", "#b45309", ["PK id", "name, unit", "stock_quantity", "min_stock", "FK branch_id"]),
        "products": (590, 650, 1035, 965, "#eff6ff", "#092379", ["PK id", "name, price", "category", "image_url", "FK branch_id", "status"]),
        "product_ingredients": (1100, 650, 1535, 930, "#ecfdf5", "#047857", ["PK id", "FK product_id", "FK ingredient_id", "quantity_needed"]),
        "stock_movements": (1615, 650, 2045, 965, "#f8fafc", "#64748b", ["PK id", "FK ingredient_id, branch_id, user_id", "movement_type, direction", "quantity", "stock_before, stock_after", "reference_type, reference_id"]),
    }
    for title, (x1, y1, x2, y2, fill, outline, lines) in boxes.items():
        draw_box(draw, (x1, y1, x2, y2), title.upper(), lines, fill, outline, font=FONT_TINY)

    draw_arrow(draw, (485, 285), (590, 285), "#092379", 3, label="1:N")
    draw_arrow(draw, (995, 285), (1100, 285), "#047857", 3, label="1:N")
    draw_arrow(draw, (1530, 285), (1615, 285), "#b91c1c", 3, label="1:N")
    draw_arrow(draw, (520, 790), (590, 790), "#b45309", 3, label="branch 1:N")
    draw_arrow(draw, (1035, 790), (1100, 790), "#047857", 3, label="1:N")
    draw_arrow(draw, (1535, 790), (1615, 790), "#64748b", 3, label="audit")
    draw_arrow(draw, (282, 395), (282, 650), "#b45309", 3, label="1:N")
    draw_arrow(draw, (792, 455), (792, 650), "#092379", 3, label="branch 1:N")

    draw_box(
        draw,
        (430, 1050, 1670, 1265),
        "RINGKASAN RELASI",
        [
            "branches 1:N users, products, ingredients, transactions, stock_movements",
            "users 1:N transactions dan users 1:N review melalui approved_by",
            "products 1:N product_ingredients dan products 1:N transaction_items",
            "ingredients 1:N product_ingredients dan ingredients 1:N stock_movements",
            "transactions 1:N transaction_items; void menghasilkan stock_movements direction in",
        ],
        "#ffffff",
        "#64748b",
        font=FONT_TINY,
    )

    img.save(ASSET_DIR / "diagram-data-model.png")


FIGURES = [
    ("Gambar 1", "Diagram use case LocalesPro sesuai fitur aplikasi saat ini.", "SKPL-assets/diagram-use-case.png"),
    ("Gambar 2", "Diagram arsitektur implementasi LocalesPro.", "SKPL-assets/diagram-architecture.png"),
    ("Gambar 3", "Sequence diagram transaksi POS tunai.", "SKPL-assets/diagram-sequence-pos-cash.png"),
    ("Gambar 4", "Data model/ERD ringkas LocalesPro.", "SKPL-assets/diagram-data-model.png"),
]

ROLE_ROWS = [
    ["Admin", "Pemilik/pengelola utama sistem.", "Dashboard, POS, cabang, akun kasir, produk, bahan baku, stok, laporan, void transaksi.", "Dapat memilih cabang aktif dan meninjau semua akun kasir."],
    ["Kasir", "Petugas penjualan harian.", "POS, laporan, stok baca-saja, riwayat mutasi cabang, dan pendaftaran akun kasir.", "Cabang dikunci berdasarkan branch_id pada akun."],
    ["Pembeli", "Pelanggan yang melakukan pembelian di outlet.", "Memilih pesanan, membayar transaksi, dan menerima struk.", "Tidak memiliki akun login di sistem."],
]

FUNCTIONAL_ROWS = [
    ["SKPL-F-001", "Autentikasi", "Sistem menyediakan login mode Kasir dan mode Admin dengan username dan password.", "Tinggi", "Pengguna valid diarahkan ke dashboard atau POS sesuai role."],
    ["SKPL-F-002", "Autentikasi", "Sistem menyimpan sesi login memakai JWT dan data user di localStorage.", "Tinggi", "Request API setelah login membawa header Authorization Bearer."],
    ["SKPL-F-003", "Otorisasi", "Sistem membatasi route berdasarkan role.", "Tinggi", "User tanpa akses diarahkan ke halaman yang sesuai rolenya."],
    ["SKPL-F-004", "Keamanan login", "Sistem menerapkan rate limit 5 percobaan gagal dalam 15 menit.", "Tinggi", "Percobaan berlebih mendapat HTTP 429 dan pesan tunggu."],
    ["SKPL-F-005", "Registrasi kasir", "Calon kasir dapat mengirim pendaftaran dengan nama, email, telepon, username, password, cabang, dan catatan.", "Tinggi", "Akun tersimpan sebagai cashier dengan status pending."],
    ["SKPL-F-006", "Review kasir", "Admin dapat menyetujui, menolak, mengaktifkan ulang, atau menonaktifkan akun kasir.", "Tinggi", "Status user berubah dan catatan review tersimpan."],
    ["SKPL-F-007", "Dashboard", "Admin melihat pendapatan hari ini, jumlah transaksi, total menu, stok menipis, dan produk terlaris.", "Sedang", "Data dihitung dari transaksi Paid, produk, dan bahan baku."],
    ["SKPL-F-008", "Cabang", "Admin dapat menambah, mengubah, menghapus, dan memilih cabang aktif.", "Tinggi", "Daftar cabang diperbarui dan pilihan cabang tersimpan."],
    ["SKPL-F-009", "Cabang", "Kasir hanya melihat dan memakai cabangnya sendiri.", "Tinggi", "Akses lintas cabang ditolak oleh resolver branch backend."],
    ["SKPL-F-010", "Produk", "Admin dapat mengelola menu, harga, kategori, status, gambar, dan resep bahan baku.", "Tinggi", "Produk tersimpan dengan minimal satu bahan resep."],
    ["SKPL-F-011", "Bahan baku", "Admin dapat mengelola nama bahan, satuan, stok, batas minimum, dan cabang.", "Tinggi", "Bahan tampil pada daftar dan dipakai untuk resep."],
    ["SKPL-F-012", "Stok", "Admin dapat menambah stok bahan dari halaman stok.", "Tinggi", "Perubahan stok menghasilkan catatan stock_movements."],
    ["SKPL-F-013", "Stok", "Kasir dapat melihat stok dan riwayat mutasi tanpa tombol update stok.", "Sedang", "Tampilan kasir bersifat baca-saja."],
    ["SKPL-F-014", "POS", "Kasir dapat memilih produk berdasarkan kategori dan memasukkannya ke keranjang.", "Tinggi", "Keranjang menampilkan item, qty, harga, dan subtotal."],
    ["SKPL-F-015", "POS", "Frontend mencegah qty melebihi stok bahan berdasarkan resep produk.", "Tinggi", "Produk habis ditandai Habis dan tidak dapat dipilih."],
    ["SKPL-F-016", "Pembayaran", "Sistem memproses pembayaran tunai/manual dan menghitung kembalian di UI.", "Tinggi", "Jika uang kurang, transaksi tidak diproses."],
    ["SKPL-F-017", "Transaksi", "Backend menyimpan transaction header, transaction_items, kode transaksi, dan status Paid.", "Tinggi", "Transaksi sukses tersimpan dalam satu DB transaction."],
    ["SKPL-F-018", "Inventori", "Sistem mengurangi stok bahan berdasarkan resep setelah transaksi sukses.", "Tinggi", "Stock movement bertipe sale tercatat."],
    ["SKPL-F-019", "Struk", "Sistem membuka jendela struk dan mencetak struk setelah transaksi berhasil.", "Sedang", "Struk memuat kode, waktu, item, total, tunai, dan kembalian."],
    ["SKPL-F-020", "Laporan", "Pengguna dapat melihat laporan transaksi, filter tanggal/metode, total pendapatan, dan mencetak laporan.", "Tinggi", "Data filter sesuai cabang aktif."],
    ["SKPL-F-021", "Void", "Pengguna dapat melakukan void terhadap transaksi Paid.", "Tinggi", "Status berubah menjadi Voided dan stok bahan dikembalikan."],
    ["SKPL-F-022", "API error", "Frontend menampilkan pesan saat API tidak dapat dihubungi.", "Sedang", "Pengguna mendapat instruksi menjalankan Apache dan MySQL XAMPP."],
]

NONFUNCTIONAL_ROWS = [
    ["Keamanan", "JWT HS256 dengan TTL 8 jam, role guard frontend, requireAuth/requireRoles backend, password_hash, dan login throttling.", "Pengguna tidak sah tidak dapat memanggil endpoint privat."],
    ["Integritas data", "Transaksi penjualan, void, update bahan, dan update produk memakai beginTransaction/commit/rollback.", "Data tidak setengah tersimpan ketika operasi gagal."],
    ["Auditabilitas", "Login dan review akun ditulis ke auth_security.log; mutasi stok ditulis ke stock_movements.", "Admin dapat menelusuri aktivitas penting."],
    ["Usability", "UI memakai bahasa Indonesia, modal konfirmasi, filter, ringkasan data, dan status badge.", "Operasi harian kasir dan admin mudah dipahami."],
    ["Kompatibilitas", "Aplikasi berjalan di browser modern, frontend Vite/React, backend PHP, dan MySQL/MariaDB via XAMPP.", "Sistem dapat dijalankan lokal di Windows/XAMPP."],
    ["Maintainability", "Kode dipisah per modul: pages React, AppContext, services/api, endpoint PHP per domain.", "Perubahan fitur dapat dilakukan per modul."],
    ["Performa", "Frontend mengambil data utama paralel dan riwayat mutasi dibatasi default 50 maksimum 200.", "Halaman tetap responsif pada data operasional normal."],
    ["Ketersediaan", "Aplikasi bergantung pada Apache dan MySQL lokal.", "Jika layanan XAMPP mati, frontend menampilkan pesan koneksi gagal."],
]

USE_CASE_ROWS = [
    ["UC-01", "Login", "Admin, Kasir", "Akun active tersedia.", "Pengguna memilih mode, mengisi credential, sistem memvalidasi role dan status, lalu membuat sesi.", "Credential salah, role tidak cocok, akun pending/rejected/inactive, atau rate limit.", "Pengguna masuk ke Dashboard atau POS."],
    ["UC-02", "Registrasi Kasir", "Kasir", "Cabang aktif tersedia.", "Kasir mengisi data, sistem validasi format, cek duplikasi, lalu menyimpan pending.", "Email/username duplikat, password lemah, cabang tidak aktif.", "Permintaan akun menunggu review admin."],
    ["UC-03", "Review Akun Kasir", "Admin", "Admin login.", "Admin membuka daftar akun, memilih cabang penempatan, memberi catatan, lalu approve/reject/activate/deactivate.", "Cabang tidak dipilih atau tidak aktif.", "Status akun berubah."],
    ["UC-04", "Kelola Cabang", "Admin", "Admin login.", "Admin tambah/edit/hapus cabang dan memilih cabang aktif.", "Data nama/alamat kosong atau cabang aktif tidak dapat dihapus lewat UI.", "Data cabang diperbarui."],
    ["UC-05", "Kelola Produk dan Resep", "Admin", "Cabang dan bahan tersedia.", "Admin mengisi nama, harga, kategori, gambar, dan resep minimal satu bahan.", "Resep kosong atau jumlah bahan tidak valid.", "Produk tersimpan dan tampil di POS."],
    ["UC-06", "Kelola Bahan Baku", "Admin", "Cabang aktif dipilih.", "Admin tambah/edit/hapus bahan, stok, satuan, dan min_stock.", "ID bahan tidak valid atau akses cabang ditolak.", "Master bahan diperbarui."],
    ["UC-07", "Update Stok", "Admin", "Bahan tersedia.", "Admin memilih bahan, mengisi jumlah masuk, sistem update stok dan catat mutasi.", "Jumlah kosong atau <= 0.", "Stok bertambah dan riwayat mutasi muncul."],
    ["UC-08", "POS Tunai", "Admin, Kasir, Pembeli", "Produk aktif tersedia dan stok resep cukup.", "Kasir/admin memilih produk pesanan pembeli, ubah qty, buka pembayaran, input uang, konfirmasi.", "Keranjang kosong, uang kurang, stok tidak cukup.", "Transaksi Paid tersimpan."],
    ["UC-09", "Cetak Struk", "Kasir, Pembeli", "Transaksi berhasil.", "Sistem membuka window struk dan menjalankan print untuk pembeli.", "Browser memblokir popup/print.", "Struk dapat dicetak atau window ditutup otomatis."],
    ["UC-10", "Laporan Transaksi", "Admin, Kasir", "Pengguna login.", "Pengguna melihat riwayat, filter tanggal/metode, melihat summary, dan print laporan.", "Tidak ada transaksi sesuai filter.", "Data laporan tampil sesuai cabang."],
    ["UC-11", "Void Transaksi", "Admin, Kasir", "Transaksi berstatus Paid.", "Pengguna klik Void, konfirmasi, backend ubah status dan restore stok.", "ID tidak valid, transaksi tidak ditemukan, sudah Voided, lintas cabang.", "Transaksi Voided dan stok kembali."],
    ["UC-12", "Lihat Riwayat Mutasi", "Admin, Kasir", "Pengguna login.", "Sistem menampilkan mutasi stok cabang aktif.", "Belum ada mutasi.", "Mutasi sale, void_restore, stock_in, stock_out terlihat."],
]

API_ROWS = [
    ["POST", "/auth/login.php", "Publik", "username, password, requested_role", "JWT, user, expires_in atau error 401/403/429"],
    ["POST", "/auth/register_cashier.php", "Publik", "full_name, email, phone, username, password, branch_id, note", "Pendaftaran cashier pending"],
    ["GET", "/public/branches_registration.php", "Publik", "-", "Daftar cabang active untuk registrasi"],
    ["GET", "/users/registrations.php", "Admin", "status", "Daftar akun kasir dan summary count"],
    ["POST", "/users/review.php", "Admin", "id, action, branch_id, review_note", "Status kasir berubah"],
    ["GET", "/branches/list.php", "Login", "branch_id implisit untuk kasir", "Daftar cabang; kasir hanya cabangnya"],
    ["POST", "/branches/create.php", "Admin", "name, address, phone, status", "Cabang baru"],
    ["POST", "/branches/update.php", "Admin", "id, name, address, phone, status", "Cabang diperbarui"],
    ["POST", "/branches/delete.php", "Admin", "id", "Cabang dihapus jika valid"],
    ["GET", "/products/read.php", "Login", "branch_id optional", "Produk active beserta recipe"],
    ["POST", "/products/create.php", "Admin", "name, price, category, image_url, status, recipe", "Produk dan recipe tersimpan"],
    ["POST", "/products/update.php", "Admin", "id dan payload produk", "Produk dan recipe diperbarui"],
    ["POST", "/products/delete.php", "Admin", "id", "Produk dihapus"],
    ["GET", "/ingredients/list.php", "Login", "branch_id optional", "Daftar bahan cabang"],
    ["POST", "/ingredients/create.php", "Admin", "name, stock, unit, minStock, branch_id", "Bahan baru"],
    ["POST", "/ingredients/update.php", "Admin", "id, name, stock, unit, minStock, branch_id", "Bahan dan stock_movement diperbarui"],
    ["POST", "/ingredients/delete.php", "Admin", "id", "Bahan dihapus"],
    ["GET", "/stock_movements/list.php", "Login", "branch_id, ingredient_id, limit", "Riwayat mutasi stok"],
    ["GET", "/transactions/history.php", "Login", "branch_id optional", "Riwayat transaksi beserta item"],
    ["POST", "/transactions/create.php", "Login", "branch_id, payment_method, items", "Transaksi Paid, transaction_items, stok terpotong"],
    ["POST", "/transactions/delete.php", "Login", "id", "Status transaksi menjadi Voided dan stok direstore"],
]

ENTITY_ROWS = [
    ["branches", "id, name, address, phone, status, created_at", "Master cabang operasional. status active/inactive."],
    ["users", "id, username, full_name, email, phone, password, role, status, registration_note, review_note, branch_id, approved_by, approved_at, last_login_at, created_at", "Akun admin, cashier dan data review kasir."],
    ["products", "id, name, price, category, image_url, branch_id, status", "Menu yang dijual di POS. Hanya status active yang tampil."],
    ["ingredients", "id, name, unit, stock_quantity, min_stock, branch_id", "Bahan baku dan stok per cabang."],
    ["product_ingredients", "id, product_id, ingredient_id, quantity_needed", "Resep produk, yaitu kebutuhan bahan per 1 produk."],
    ["transactions", "id, transaction_code, user_id, branch_id, total_price, payment_method, payment_gateway, amount_paid, change_amount, payment_status, payment_note, created_at, paid_at", "Header transaksi penjualan tunai/manual."],
    ["transaction_items", "id, transaction_id, product_id, quantity, subtotal", "Detail produk dalam satu transaksi."],
    ["stock_movements", "id, ingredient_id, branch_id, user_id, movement_type, direction, quantity, stock_before, stock_after, reference_type, reference_id, notes, created_at", "Riwayat mutasi stok manual, sale, dan void_restore."],
]

BUSINESS_RULE_ROWS = [
    ["RB-01", "Role login harus sesuai mode: cashier untuk mode kasir dan admin untuk mode dashboard."],
    ["RB-02", "Akun dengan status pending, rejected, atau inactive tidak boleh login."],
    ["RB-03", "Kasir hanya boleh mengakses data cabangnya sendiri."],
    ["RB-04", "Admin wajib memilih cabang aktif untuk operasi data cabang tertentu."],
    ["RB-05", "Produk yang dapat dijual harus berstatus active dan memiliki resep minimal satu bahan."],
    ["RB-06", "Frontend POS menghitung ketersediaan produk dari stok bahan dibagi kebutuhan resep."],
    ["RB-07", "Pembayaran yang didukung aplikasi saat ini adalah Cash/manual."],
    ["RB-08", "Transaksi sukses langsung berstatus Paid dan diberi kode LOC-YYYYMMDD-xxxxx."],
    ["RB-09", "Setelah transaksi Paid tersimpan, stok bahan berkurang sesuai resep dan qty item."],
    ["RB-10", "Void hanya memulihkan stok jika status sebelumnya Paid dan status baru Voided."],
    ["RB-11", "Catatan customer/payment pada modal POS digunakan pada struk UI; penyimpanan backend saat ini fokus pada data transaksi inti."],
    ["RB-12", "Jika API dipanggil langsung, backend memvalidasi token, role, cabang, produk aktif, dan item; validasi stok negatif masih bergantung pada alur UI POS."],
]

WHITEBOX_ROWS = [
    ["WB-01", "frontend/src/pages/auth/Login.jsx", "validateLoginInput: username < 4", "user=abc, password valid", "Mengembalikan pesan Username minimal 4 karakter.", "Valid dari kode"],
    ["WB-02", "frontend/src/pages/auth/Login.jsx", "validateLoginInput: username karakter ilegal", "user=kasir 1", "Mengembalikan pesan pola username.", "Valid dari kode"],
    ["WB-03", "backend/api/auth/login.php", "requested_role tidak dalam cashier/dashboard", "requested_role=owner", "HTTP 422 Mode login tidak valid.", "Valid dari kode"],
    ["WB-04", "backend/api/auth/login.php", "role mismatch dashboard dengan cashier", "cashier login mode dashboard", "HTTP 403 tidak punya akses mode admin.", "Valid dari kode"],
    ["WB-05", "backend/api/auth/login.php", "5 kali credential salah", "password salah berulang", "Percobaan berikutnya HTTP 429.", "Valid dari kode"],
    ["WB-06", "backend/api/auth/register_cashier.php", "password kurang kuat", "password tanpa huruf besar/angka", "HTTP 422 dengan pesan strength password.", "Valid dari kode"],
    ["WB-07", "backend/api/auth/register_cashier.php", "duplikasi username/email", "username/email sudah ada", "HTTP 409 username/email sudah dipakai.", "Valid dari kode"],
    ["WB-08", "backend/api/users/review.php", "approve tanpa branch valid", "action=approve, branch_id=0", "HTTP 422 Cabang penempatan wajib dipilih.", "Valid dari kode"],
    ["WB-09", "frontend/src/routes/AppRoutes.jsx", "ProtectedRoute tanpa session", "localStorage kosong", "Redirect ke /login.", "Valid dari kode"],
    ["WB-10", "frontend/src/routes/AppRoutes.jsx", "Cashier membuka dashboard", "role=cashier path=/dashboard", "Redirect ke /pos.", "Valid dari kode"],
    ["WB-11", "frontend/src/store/AppContext.jsx", "syncSelectedBranch branch tersimpan hilang", "selectedBranch id tidak ada di list", "Set selectedBranch ke branch pertama.", "Valid dari kode"],
    ["WB-12", "backend/api/products/create.php", "recipe kosong", "recipe=[]", "HTTP 400 Resep produk wajib diisi minimal 1 bahan.", "Valid dari kode"],
    ["WB-13", "backend/api/products/update.php", "quantity_needed <= 0", "amount=0", "HTTP 400 Jumlah bahan harus lebih dari 0.", "Valid dari kode"],
    ["WB-14", "frontend/src/pages/pos/POS.jsx", "getAvailableStockCount tanpa recipe", "product.recipe kosong", "Mengembalikan Infinity; produk tidak dianggap habis karena resep belum dipakai.", "Valid dari kode"],
    ["WB-15", "frontend/src/pages/pos/POS.jsx", "addToCart qty melewati maxAvailable", "qty cart >= stok bahan", "Menampilkan alert stok tidak cukup dan qty tidak bertambah.", "Valid dari kode"],
    ["WB-16", "frontend/src/pages/pos/POS.jsx", "handleCashPayment uang kurang", "cashAmount < totalPrice", "paymentError uang diterima kurang; tidak POST transaksi.", "Valid dari kode"],
    ["WB-17", "backend/api/transactions/create.php", "items kosong", "items=[]", "HTTP 422 Item transaksi tidak boleh kosong.", "Valid dari kode"],
    ["WB-18", "backend/api/transactions/create.php", "productId/qty tidak valid", "id=0 atau qty=0", "Rollback dan HTTP 422 item tidak valid.", "Valid dari kode"],
    ["WB-19", "backend/api/transactions/create.php", "produk inactive/tidak ada", "product id tidak tersedia", "Rollback dan HTTP 422 produk tidak tersedia.", "Valid dari kode"],
    ["WB-20", "backend/api/transactions/create.php", "produk berbeda cabang", "product.branch_id != branch aktif", "Rollback dan HTTP 422 produk tidak tersedia untuk cabang aktif.", "Valid dari kode"],
    ["WB-21", "backend/api/transactions/create.php", "path sukses", "items valid dan branch valid", "Insert transactions, update kode, insert items, deduct inventory, commit.", "Valid dari kode"],
    ["WB-22", "backend/api/transactions/delete.php", "transaction id <= 0", "id=0", "HTTP 422 ID transaksi tidak valid.", "Valid dari kode"],
    ["WB-23", "backend/api/transactions/delete.php", "transaction sudah Voided", "payment_status=Voided", "Tidak restore ulang; pesan transaksi sudah void.", "Valid dari kode"],
    ["WB-24", "backend/api/transactions/delete.php", "void Paid", "payment_status=Paid", "Update Voided, apply restore, commit.", "Valid dari kode"],
    ["WB-25", "backend/api/ingredients/update.php", "stockDelta positif", "stock lama 10, baru 15", "Record stock_movement type stock_in direction in quantity 5.", "Valid dari kode"],
    ["WB-26", "backend/api/stock_movements/list.php", "limit di luar batas", "limit=1000", "Limit dibatasi maksimum 200.", "Valid dari kode"],
]

BLACKBOX_ROWS = [
    ["BB-01", "Login Admin", "Masuk dengan akun admin aktif pada mode Admin.", "username/password admin valid", "Masuk ke /dashboard dan sidebar admin tampil.", "Diharapkan berhasil"],
    ["BB-02", "Login Kasir", "Masuk dengan akun kasir aktif pada mode Kasir.", "username/password kasir valid", "Masuk ke /pos dan cabang terkunci.", "Diharapkan berhasil"],
    ["BB-03", "Login invalid", "Credential salah.", "username valid password salah", "Pesan login gagal dan sisa percobaan tampil.", "Diharapkan berhasil"],
    ["BB-04", "Login mode salah", "Kasir mencoba mode Admin.", "akun kasir, mode dashboard", "Akses ditolak.", "Diharapkan berhasil"],
    ["BB-05", "Registrasi kasir", "Kirim pendaftaran valid.", "Nama 2 kata, email valid, HP valid, password kuat, cabang aktif", "Pesan pendaftaran terkirim dan akun pending.", "Diharapkan berhasil"],
    ["BB-06", "Registrasi invalid email", "Isi email tanpa format email.", "email=abc", "Pesan Email tidak valid.", "Diharapkan berhasil"],
    ["BB-07", "Registrasi password lemah", "Password kurang dari aturan.", "password=abcdefg", "Pesan password minimal/kompleksitas.", "Diharapkan berhasil"],
    ["BB-08", "Review approve", "Admin menyetujui akun pending.", "action approve + branch aktif", "Akun menjadi active.", "Diharapkan berhasil"],
    ["BB-09", "Review reject", "Admin menolak akun pending.", "action reject + catatan", "Akun menjadi rejected.", "Diharapkan berhasil"],
    ["BB-10", "Login akun pending", "Akun kasir pending mencoba login.", "credential pending", "Login ditolak menunggu persetujuan admin.", "Diharapkan berhasil"],
    ["BB-11", "Dashboard", "Admin membuka dashboard.", "data transaksi dan stok tersedia", "Kartu pendapatan, pesanan, menu, stok menipis, produk terlaris tampil.", "Diharapkan berhasil"],
    ["BB-12", "Tambah cabang", "Admin menambah cabang valid.", "name, address, phone, status", "Cabang baru tampil pada daftar.", "Diharapkan berhasil"],
    ["BB-13", "Tambah cabang tidak lengkap", "Admin submit tanpa nama/alamat.", "name kosong atau address kosong", "Pesan data tidak lengkap.", "Diharapkan berhasil"],
    ["BB-14", "Pilih cabang", "Admin memilih cabang operasional.", "klik Pilih Cabang Ini", "Header lokasi berubah ke cabang dipilih.", "Diharapkan berhasil"],
    ["BB-15", "Tambah bahan", "Admin tambah bahan.", "name, stock, unit, minStock", "Bahan tampil pada tabel bahan baku.", "Diharapkan berhasil"],
    ["BB-16", "Update stok", "Admin tambah stok bahan.", "jumlah masuk > 0", "Stok bertambah dan riwayat mutasi bertipe Tambah Stok muncul.", "Diharapkan berhasil"],
    ["BB-17", "Update stok invalid", "Submit jumlah kosong/0.", "amount=0", "Alert jumlah stok tidak valid.", "Diharapkan berhasil"],
    ["BB-18", "Kasir lihat stok", "Kasir membuka halaman stok.", "role cashier", "Tombol update tidak tampil; label Lihat Saja tampil.", "Diharapkan berhasil"],
    ["BB-19", "Tambah produk", "Admin tambah menu dengan resep.", "name, price, category, image, recipe", "Produk tampil pada Manajemen Menu dan POS.", "Diharapkan berhasil"],
    ["BB-20", "Produk tanpa resep", "Submit produk tanpa bahan.", "recipe kosong", "Alert/response resep wajib diisi.", "Diharapkan berhasil"],
    ["BB-21", "Lihat resep", "Pengguna klik Lihat Resep.", "produk memiliki recipe", "Modal resep menampilkan bahan, jumlah, stok tersedia, status cukup/kurang.", "Diharapkan berhasil"],
    ["BB-22", "POS tambah item", "Kasir memilih produk aktif.", "klik kartu produk", "Item masuk keranjang dan total berubah.", "Diharapkan berhasil"],
    ["BB-23", "POS stok habis", "Produk dengan stok resep 0.", "klik produk habis", "Produk disabled atau alert stok tidak cukup.", "Diharapkan berhasil"],
    ["BB-24", "Pembayaran uang kurang", "Kasir mengisi uang kurang dari total.", "cashReceived < total", "Pesan uang diterima kurang; transaksi tidak tersimpan.", "Diharapkan berhasil"],
    ["BB-25", "Pembayaran tunai valid", "Kasir membayar dengan uang cukup.", "cashReceived >= total", "Transaksi berhasil, struk terbuka, keranjang kosong.", "Diharapkan berhasil"],
    ["BB-26", "Riwayat transaksi", "Buka laporan setelah transaksi.", "cabang aktif", "Transaksi baru muncul dengan status Paid.", "Diharapkan berhasil"],
    ["BB-27", "Filter laporan", "Filter berdasarkan tanggal/metode.", "tanggal tertentu, metode Cash", "Tabel hanya menampilkan transaksi sesuai filter dan summary ikut berubah.", "Diharapkan berhasil"],
    ["BB-28", "Void transaksi", "Void transaksi Paid.", "klik Void lalu konfirmasi", "Status berubah Voided dan tombol void hilang.", "Diharapkan berhasil"],
    ["BB-29", "Print laporan", "Klik Cetak Laporan.", "data laporan tampil", "Dialog print browser dipanggil.", "Diharapkan berhasil"],
    ["BB-30", "API mati", "Frontend membuka data saat API tidak bisa dihubungi.", "Apache/MySQL mati", "Pesan koneksi API gagal tampil.", "Diharapkan berhasil"],
]

TRACE_ROWS = [
    ["SKPL-F-001 s.d. F-004", "UC-01", "WB-01 s.d. WB-05", "BB-01 s.d. BB-04"],
    ["SKPL-F-005", "UC-02", "WB-06, WB-07", "BB-05 s.d. BB-07, BB-10"],
    ["SKPL-F-006", "UC-03", "WB-08", "BB-08, BB-09"],
    ["SKPL-F-007", "UC-01, UC-10", "WB-09, WB-10", "BB-11"],
    ["SKPL-F-008, F-009", "UC-04", "WB-11", "BB-12 s.d. BB-14"],
    ["SKPL-F-010", "UC-05", "WB-12, WB-13", "BB-19 s.d. BB-21"],
    ["SKPL-F-011 s.d. F-013", "UC-06, UC-07, UC-12", "WB-25, WB-26", "BB-15 s.d. BB-18"],
    ["SKPL-F-014 s.d. F-019", "UC-08, UC-09", "WB-14 s.d. WB-21", "BB-22 s.d. BB-26"],
    ["SKPL-F-020, F-021", "UC-10, UC-11", "WB-22 s.d. WB-24", "BB-27 s.d. BB-29"],
    ["SKPL-F-022", "Semua use case API", "Interceptor frontend", "BB-30"],
]


def md_escape(value):
    return str(value).replace("|", "\\|").replace("\n", "<br>")


def md_table(headers, rows):
    lines = [
        "| " + " | ".join(md_escape(header) for header in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(md_escape(cell) for cell in row) + " |")
    return lines


def build_markdown():
    lines = []
    add = lines.append
    extend = lines.extend

    add("# SKPL Locales Pro")
    add("")
    add("**Spesifikasi Kebutuhan Perangkat Lunak (SKPL)**")
    add("")
    add(f"**Versi dokumen:** 2.0  ")
    add(f"**Tanggal pemutakhiran:** {REPORT_DATE}  ")
    add("**Objek:** Aplikasi Point of Sale Locales Pro  ")
    add("**Teknologi:** React/Vite, PHP API, MySQL/MariaDB, XAMPP")
    add("")
    add("Dokumen ini disusun ulang berdasarkan struktur aplikasi Locales Pro yang ada di repository saat ini. Fokus aplikasi saat ini adalah POS minuman/cafe dengan transaksi tunai/manual, manajemen cabang, produk, resep, bahan baku, stok, laporan transaksi, void transaksi, dan review akun kasir.")
    add("")
    add("## Riwayat Revisi")
    extend(md_table(["Versi", "Tanggal", "Keterangan"], [["1.0", "Sebelumnya", "Dokumen ringkas awal."], ["2.0", REPORT_DATE, "Dilengkapi sesuai aplikasi saat ini, termasuk diagram, kebutuhan fungsional, data model, whitebox, dan blackbox."]]))
    add("")
    add("## Daftar Gambar")
    extend(md_table(["No", "Nama Gambar", "File"], FIGURES))
    add("")
    add("## Daftar Tabel")
    table_list = [
        ["Tabel 1", "Riwayat revisi"],
        ["Tabel 2", "Daftar gambar"],
        ["Tabel 3", "Daftar tabel"],
        ["Tabel 4", "Aktor dan hak akses"],
        ["Tabel 5", "Kebutuhan fungsional"],
        ["Tabel 6", "Kebutuhan non-fungsional"],
        ["Tabel 7", "Rincian use case"],
        ["Tabel 8", "Endpoint API"],
        ["Tabel 9", "Kamus data entitas"],
        ["Tabel 10", "Aturan bisnis"],
        ["Tabel 11", "Whitebox testing"],
        ["Tabel 12", "Blackbox testing"],
        ["Tabel 13", "Matriks traceability"],
    ]
    extend(md_table(["No", "Nama Tabel"], table_list))
    add("")

    add("## 1. Pendahuluan")
    add("")
    add("### 1.1 Tujuan")
    add("SKPL ini menjelaskan kebutuhan, batasan, alur kerja, data, antarmuka, dan rancangan pengujian aplikasi Locales Pro. Dokumen dipakai sebagai acuan pengembangan, pengujian, dan penilaian aplikasi.")
    add("")
    add("### 1.2 Ruang Lingkup")
    add("Ruang lingkup sistem meliputi aplikasi POS lokal untuk operasional minuman/cafe. Sistem mendukung tiga aktor utama: admin, kasir, dan pembeli. Integrasi payment gateway eksternal tidak dipakai pada implementasi saat ini; pembayaran yang aktif adalah tunai/manual.")
    add("")
    add("### 1.3 Definisi dan Singkatan")
    extend(md_table(["Istilah", "Definisi"], [
        ["SKPL", "Spesifikasi Kebutuhan Perangkat Lunak."],
        ["POS", "Point of Sale, halaman kasir untuk transaksi penjualan."],
        ["JWT", "JSON Web Token untuk sesi login."],
        ["CRUD", "Create, Read, Update, Delete."],
        ["Recipe/Resep", "Daftar bahan dan jumlah yang dibutuhkan untuk membuat satu produk."],
        ["Void", "Pembatalan transaksi yang mengubah status menjadi Voided dan mengembalikan stok."],
        ["Stock movement", "Catatan mutasi stok masuk/keluar."],
    ]))
    add("")
    add("### 1.4 Referensi Implementasi")
    extend(md_table(["Area", "File/Folder"], [
        ["Frontend route dan role", "frontend/src/routes/AppRoutes.jsx"],
        ["State aplikasi", "frontend/src/store/AppContext.jsx"],
        ["API client", "frontend/src/services/api.js"],
        ["Autentikasi backend", "backend/auth/verify_token.php dan backend/api/auth/login.php"],
        ["Transaksi dan void", "backend/api/transactions/create.php dan backend/api/transactions/delete.php"],
        ["Inventori", "backend/config/inventory_helpers.php dan backend/config/payment_helpers.php"],
        ["Database utama", "DB/locales_db.sql dan backend/database/*.sql"],
    ]))
    add("")

    add("## 2. Gambaran Umum Sistem")
    add("")
    add("Locales Pro adalah aplikasi POS berbasis web yang berjalan lokal di XAMPP. Frontend React/Vite berkomunikasi dengan backend PHP melalui JSON API. Backend memakai PDO ke MySQL/MariaDB. Data sesi tersimpan di browser localStorage dan dikirim sebagai Bearer token pada request API.")
    add("")
    add("### 2.1 Gambar Sistem")
    for label, caption, path in FIGURES:
        add(f"![{label}. {caption}]({path})")
        add("")
    add("### 2.2 Aktor dan Hak Akses")
    extend(md_table(["Aktor", "Deskripsi", "Hak Akses", "Catatan"], ROLE_ROWS))
    add("")
    add("### 2.3 Lingkungan Operasi")
    extend(md_table(["Komponen", "Spesifikasi"], [
        ["Frontend", "React 19, Vite 8, react-router-dom, axios, lucide-react, chart.js."],
        ["Backend", "PHP API modular dengan konfigurasi CORS dan PDO MySQL."],
        ["Database", "MySQL/MariaDB database locales_db."],
        ["Runtime lokal", "XAMPP Apache dan MySQL."],
        ["Base API default", "http://localhost/LocalesPro-v1-main/backend/api"],
    ]))
    add("")

    add("## 3. Kebutuhan Fungsional")
    extend(md_table(["ID", "Modul", "Kebutuhan", "Prioritas", "Kriteria Penerimaan"], FUNCTIONAL_ROWS))
    add("")
    add("## 4. Kebutuhan Non-Fungsional")
    extend(md_table(["Kategori", "Kebutuhan", "Kriteria Penerimaan"], NONFUNCTIONAL_ROWS))
    add("")

    add("## 5. Use Case")
    add("Use case berikut menjelaskan perilaku sistem dari sisi aktor. Alur yang dicantumkan mengikuti halaman dan endpoint yang tersedia pada aplikasi saat ini.")
    add("")
    extend(md_table(["ID", "Nama", "Aktor", "Prasyarat", "Alur Utama", "Alternatif/Error", "Hasil Akhir"], USE_CASE_ROWS))
    add("")

    add("## 6. Antarmuka Eksternal")
    add("")
    add("### 6.1 Antarmuka Pengguna")
    extend(md_table(["Halaman", "Isi Antarmuka", "Role"], [
        ["Login", "Mode Kasir dan Admin, input username/password, toggle lihat password, link daftar kasir.", "Publik"],
        ["Registrasi Kasir", "Form identitas, cabang tujuan, password, catatan pendaftaran.", "Publik"],
        ["Dashboard", "Kartu pendapatan, pesanan, total menu, stok menipis, produk terlaris.", "Admin"],
        ["Cabang", "Kartu cabang, modal tambah/edit, pilihan cabang aktif.", "Admin"],
        ["Akun Kasir", "Filter status, kartu akun, select cabang, catatan review, aksi approve/reject/activate/deactivate.", "Admin"],
        ["Produk", "Tabel menu, gambar, resep, modal tambah/edit, modal lihat resep.", "Admin"],
        ["Bahan Baku", "Tabel bahan, stok, satuan, min_stock, modal tambah/edit.", "Admin"],
        ["Stok", "Kartu stok bahan, modal update stok untuk admin, riwayat mutasi.", "Admin, Kasir baca-saja"],
        ["POS", "Daftar produk, filter kategori, keranjang, modal pembayaran tunai, cetak struk.", "Admin, Kasir, Pembeli"],
        ["Laporan", "Summary, filter tanggal/metode, tabel transaksi, void, print laporan.", "Admin, Kasir"],
    ]))
    add("")
    add("### 6.2 Endpoint API")
    extend(md_table(["Method", "Endpoint", "Akses", "Input", "Output"], API_ROWS))
    add("")

    add("## 7. Kebutuhan Data")
    add("")
    add("### 7.1 Kamus Data")
    extend(md_table(["Entitas", "Kolom Penting", "Keterangan"], ENTITY_ROWS))
    add("")
    add("### 7.2 Relasi Data")
    add("Relasi utama sistem: satu branch memiliki banyak users, products, ingredients, transactions, dan stock_movements. Satu product memiliki banyak product_ingredients dan dapat muncul di banyak transaction_items. Satu transaction memiliki banyak transaction_items. Void dan penjualan menggunakan stock_movements sebagai catatan audit stok.")
    add("")

    add("## 8. Aturan Bisnis dan Batasan Implementasi")
    extend(md_table(["ID", "Aturan"], BUSINESS_RULE_ROWS))
    add("")
    add("Catatan penting: aplikasi sekarang tidak memakai Midtrans atau payment gateway lain. File konfigurasi pembayaran masih ada sebagai konfigurasi dasar, tetapi alur aktif pada UI dan API adalah Cash/manual.")
    add("")

    add("## 9. Rancangan Pengujian Whitebox")
    add("Whitebox berikut diturunkan dari cabang kondisi pada kode frontend dan backend. Status 'Valid dari kode' berarti perilaku tersebut terlihat langsung dari pembacaan source code dan tetap perlu dieksekusi ulang pada environment demo bila diminta sebagai bukti uji runtime.")
    add("")
    extend(md_table(["ID", "Modul/Path", "Kondisi Internal", "Data Uji", "Ekspektasi", "Status"], WHITEBOX_ROWS))
    add("")

    add("## 10. Rancangan Pengujian Blackbox")
    add("Blackbox berikut berfokus pada masukan dan keluaran yang terlihat pengguna tanpa melihat kode internal.")
    add("")
    extend(md_table(["ID", "Fitur", "Skenario", "Input", "Output yang Diharapkan", "Status"], BLACKBOX_ROWS))
    add("")

    add("## 11. Matriks Traceability")
    extend(md_table(["Kebutuhan", "Use Case", "Whitebox", "Blackbox"], TRACE_ROWS))
    add("")

    add("## 12. Risiko dan Rekomendasi")
    extend(md_table(["No", "Risiko/Batasan", "Rekomendasi"], [
        ["1", "Backend transaksi belum menolak stok negatif bila API dipanggil langsung tanpa UI POS.", "Tambahkan validasi stok cukup di backend sebelum insert transaksi dan deduct inventory."],
        ["2", "Cash received, customer name, dan payment note pada modal POS dipakai untuk struk UI, tetapi backend saat ini menyimpan amount_paid setara total dan change_amount 0.", "Tambahkan field amount_paid, change_amount, customer_name, dan payment_note pada payload transaksi bila data tersebut harus muncul di laporan."],
        ["3", "Update recipe memakai ON DUPLICATE KEY; skema product_ingredients saat ini belum terlihat memiliki unique key gabungan product_id dan ingredient_id.", "Tambahkan unique index product_id, ingredient_id agar upsert backend konsisten."],
        ["4", "Aplikasi bergantung pada XAMPP lokal.", "Siapkan checklist start Apache/MySQL dan backup database sebelum demo."],
    ]))
    add("")

    add("## 13. Penutup")
    add("SKPL ini mencerminkan kondisi aplikasi Locales Pro saat ini: POS tunai/manual dengan aktor admin, kasir, dan pembeli, registrasi kasir, manajemen cabang, produk, resep, bahan baku, stok, laporan, void, dan audit mutasi stok. Dokumen ini dapat dipakai sebagai dasar pengujian dan pengembangan lanjutan.")
    add("")

    REPORT_MD.write_text("\n".join(lines), encoding="utf-8")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8.5, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_doc_table(doc, headers, rows, caption=None, font_size=8.3):
    if caption:
        paragraph = doc.add_paragraph()
        paragraph.style = doc.styles["Caption"]
        paragraph.add_run(caption).bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=font_size, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "092379")
    for row in rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            set_cell_text(cells[i], cell_text, size=font_size)
    doc.add_paragraph("")
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def add_figure(doc, label, caption, path):
    image_path = ROOT / path
    doc.add_picture(str(image_path), width=Inches(9.8))
    paragraph = doc.add_paragraph(f"{label}. {caption}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.italic = True
        run.font.size = Pt(9)
    doc.add_paragraph("")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(9)
    doc.styles["Heading 1"].font.name = "Arial"
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = RGBColor(9, 35, 121)
    doc.styles["Heading 2"].font.name = "Arial"
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 2"].font.color.rgb = RGBColor(9, 35, 121)

    logo_path = ROOT / "frontend" / "src" / "assets" / "locales1.png"
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo_path), width=Inches(1.35))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SKPL LOCALES PRO")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(9, 35, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Spesifikasi Kebutuhan Perangkat Lunak").bold = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Versi 2.0 - {REPORT_DATE}\nReact/Vite, PHP API, MySQL/MariaDB, XAMPP")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Dokumen ini disusun berdasarkan aplikasi Locales Pro saat ini. Fokus implementasi adalah POS tunai/manual, manajemen cabang, produk, resep, bahan baku, stok, laporan transaksi, void transaksi, dan review akun kasir."
    )
    doc.add_page_break()

    doc.add_heading("Riwayat Revisi", level=1)
    add_doc_table(doc, ["Versi", "Tanggal", "Keterangan"], [["1.0", "Sebelumnya", "Dokumen ringkas awal."], ["2.0", REPORT_DATE, "Dilengkapi sesuai aplikasi saat ini, termasuk diagram, kebutuhan fungsional, data model, whitebox, dan blackbox."]], "Tabel 1. Riwayat revisi")

    doc.add_heading("Daftar Gambar", level=1)
    add_doc_table(doc, ["No", "Nama Gambar", "File"], FIGURES, "Tabel 2. Daftar gambar")
    doc.add_heading("Daftar Tabel", level=1)
    add_doc_table(doc, ["No", "Nama Tabel"], [
        ["Tabel 1", "Riwayat revisi"],
        ["Tabel 2", "Daftar gambar"],
        ["Tabel 3", "Daftar tabel"],
        ["Tabel 4", "Aktor dan hak akses"],
        ["Tabel 5", "Kebutuhan fungsional"],
        ["Tabel 6", "Kebutuhan non-fungsional"],
        ["Tabel 7", "Rincian use case"],
        ["Tabel 8", "Endpoint API"],
        ["Tabel 9", "Kamus data entitas"],
        ["Tabel 10", "Aturan bisnis"],
        ["Tabel 11", "Whitebox testing"],
        ["Tabel 12", "Blackbox testing"],
        ["Tabel 13", "Matriks traceability"],
    ], "Tabel 3. Daftar tabel")

    doc.add_heading("1. Pendahuluan", level=1)
    doc.add_heading("1.1 Tujuan", level=2)
    doc.add_paragraph("SKPL ini menjelaskan kebutuhan, batasan, alur kerja, data, antarmuka, dan rancangan pengujian aplikasi Locales Pro. Dokumen dipakai sebagai acuan pengembangan, pengujian, dan penilaian aplikasi.")
    doc.add_heading("1.2 Ruang Lingkup", level=2)
    doc.add_paragraph("Ruang lingkup sistem meliputi aplikasi POS lokal untuk operasional minuman/cafe. Sistem mendukung tiga aktor utama: admin, kasir, dan pembeli. Integrasi payment gateway eksternal tidak dipakai pada implementasi saat ini; pembayaran yang aktif adalah tunai/manual.")
    doc.add_heading("1.3 Definisi dan Singkatan", level=2)
    add_doc_table(doc, ["Istilah", "Definisi"], [
        ["SKPL", "Spesifikasi Kebutuhan Perangkat Lunak."],
        ["POS", "Point of Sale, halaman kasir untuk transaksi penjualan."],
        ["JWT", "JSON Web Token untuk sesi login."],
        ["CRUD", "Create, Read, Update, Delete."],
        ["Recipe/Resep", "Daftar bahan dan jumlah yang dibutuhkan untuk membuat satu produk."],
        ["Void", "Pembatalan transaksi yang mengubah status menjadi Voided dan mengembalikan stok."],
        ["Stock movement", "Catatan mutasi stok masuk/keluar."],
    ])
    doc.add_heading("1.4 Referensi Implementasi", level=2)
    add_doc_table(doc, ["Area", "File/Folder"], [
        ["Frontend route dan role", "frontend/src/routes/AppRoutes.jsx"],
        ["State aplikasi", "frontend/src/store/AppContext.jsx"],
        ["API client", "frontend/src/services/api.js"],
        ["Autentikasi backend", "backend/auth/verify_token.php dan backend/api/auth/login.php"],
        ["Transaksi dan void", "backend/api/transactions/create.php dan backend/api/transactions/delete.php"],
        ["Inventori", "backend/config/inventory_helpers.php dan backend/config/payment_helpers.php"],
        ["Database utama", "DB/locales_db.sql dan backend/database/*.sql"],
    ])

    doc.add_heading("2. Gambaran Umum Sistem", level=1)
    doc.add_paragraph("Locales Pro adalah aplikasi POS berbasis web yang berjalan lokal di XAMPP. Frontend React/Vite berkomunikasi dengan backend PHP melalui JSON API. Backend memakai PDO ke MySQL/MariaDB. Data sesi tersimpan di browser localStorage dan dikirim sebagai Bearer token pada request API.")
    for label, caption, path in FIGURES:
        add_figure(doc, label, caption, path)
    doc.add_heading("2.1 Aktor dan Hak Akses", level=2)
    add_doc_table(doc, ["Aktor", "Deskripsi", "Hak Akses", "Catatan"], ROLE_ROWS, "Tabel 4. Aktor dan hak akses", 8)
    doc.add_heading("2.2 Lingkungan Operasi", level=2)
    add_doc_table(doc, ["Komponen", "Spesifikasi"], [
        ["Frontend", "React 19, Vite 8, react-router-dom, axios, lucide-react, chart.js."],
        ["Backend", "PHP API modular dengan konfigurasi CORS dan PDO MySQL."],
        ["Database", "MySQL/MariaDB database locales_db."],
        ["Runtime lokal", "XAMPP Apache dan MySQL."],
        ["Base API default", "http://localhost/LocalesPro-v1-main/backend/api"],
    ])

    doc.add_heading("3. Kebutuhan Fungsional", level=1)
    add_doc_table(doc, ["ID", "Modul", "Kebutuhan", "Prioritas", "Kriteria Penerimaan"], FUNCTIONAL_ROWS, "Tabel 5. Kebutuhan fungsional", 7.7)

    doc.add_heading("4. Kebutuhan Non-Fungsional", level=1)
    add_doc_table(doc, ["Kategori", "Kebutuhan", "Kriteria Penerimaan"], NONFUNCTIONAL_ROWS, "Tabel 6. Kebutuhan non-fungsional", 8)

    doc.add_heading("5. Use Case", level=1)
    doc.add_paragraph("Use case berikut menjelaskan perilaku sistem dari sisi aktor. Alur yang dicantumkan mengikuti halaman dan endpoint yang tersedia pada aplikasi saat ini.")
    add_doc_table(doc, ["ID", "Nama", "Aktor", "Prasyarat", "Alur Utama", "Alternatif/Error", "Hasil Akhir"], USE_CASE_ROWS, "Tabel 7. Rincian use case", 7.1)

    doc.add_heading("6. Antarmuka Eksternal", level=1)
    doc.add_heading("6.1 Antarmuka Pengguna", level=2)
    add_doc_table(doc, ["Halaman", "Isi Antarmuka", "Role"], [
        ["Login", "Mode Kasir dan Admin, input username/password, toggle lihat password, link daftar kasir.", "Publik"],
        ["Registrasi Kasir", "Form identitas, cabang tujuan, password, catatan pendaftaran.", "Publik"],
        ["Dashboard", "Kartu pendapatan, pesanan, total menu, stok menipis, produk terlaris.", "Admin"],
        ["Cabang", "Kartu cabang, modal tambah/edit, pilihan cabang aktif.", "Admin"],
        ["Akun Kasir", "Filter status, kartu akun, select cabang, catatan review, aksi approve/reject/activate/deactivate.", "Admin"],
        ["Produk", "Tabel menu, gambar, resep, modal tambah/edit, modal lihat resep.", "Admin"],
        ["Bahan Baku", "Tabel bahan, stok, satuan, min_stock, modal tambah/edit.", "Admin"],
        ["Stok", "Kartu stok bahan, modal update stok untuk admin, riwayat mutasi.", "Admin, Kasir baca-saja"],
        ["POS", "Daftar produk, filter kategori, keranjang, modal pembayaran tunai, cetak struk.", "Admin, Kasir, Pembeli"],
        ["Laporan", "Summary, filter tanggal/metode, tabel transaksi, void, print laporan.", "Admin, Kasir"],
    ])
    doc.add_heading("6.2 Endpoint API", level=2)
    add_doc_table(doc, ["Method", "Endpoint", "Akses", "Input", "Output"], API_ROWS, "Tabel 8. Endpoint API", 7.3)

    doc.add_heading("7. Kebutuhan Data", level=1)
    add_doc_table(doc, ["Entitas", "Kolom Penting", "Keterangan"], ENTITY_ROWS, "Tabel 9. Kamus data entitas", 7.4)
    doc.add_paragraph("Relasi utama sistem: satu branch memiliki banyak users, products, ingredients, transactions, dan stock_movements. Satu product memiliki banyak product_ingredients dan dapat muncul di banyak transaction_items. Satu transaction memiliki banyak transaction_items. Void dan penjualan menggunakan stock_movements sebagai catatan audit stok.")

    doc.add_heading("8. Aturan Bisnis dan Batasan Implementasi", level=1)
    add_doc_table(doc, ["ID", "Aturan"], BUSINESS_RULE_ROWS, "Tabel 10. Aturan bisnis", 8)
    doc.add_paragraph("Catatan penting: aplikasi sekarang tidak memakai Midtrans atau payment gateway lain. File konfigurasi pembayaran masih ada sebagai konfigurasi dasar, tetapi alur aktif pada UI dan API adalah Cash/manual.")

    doc.add_heading("9. Rancangan Pengujian Whitebox", level=1)
    doc.add_paragraph("Whitebox berikut diturunkan dari cabang kondisi pada kode frontend dan backend. Status 'Valid dari kode' berarti perilaku tersebut terlihat langsung dari pembacaan source code dan tetap perlu dieksekusi ulang pada environment demo bila diminta sebagai bukti uji runtime.")
    add_doc_table(doc, ["ID", "Modul/Path", "Kondisi Internal", "Data Uji", "Ekspektasi", "Status"], WHITEBOX_ROWS, "Tabel 11. Whitebox testing", 6.8)

    doc.add_heading("10. Rancangan Pengujian Blackbox", level=1)
    doc.add_paragraph("Blackbox berikut berfokus pada masukan dan keluaran yang terlihat pengguna tanpa melihat kode internal.")
    add_doc_table(doc, ["ID", "Fitur", "Skenario", "Input", "Output yang Diharapkan", "Status"], BLACKBOX_ROWS, "Tabel 12. Blackbox testing", 6.9)

    doc.add_heading("11. Matriks Traceability", level=1)
    add_doc_table(doc, ["Kebutuhan", "Use Case", "Whitebox", "Blackbox"], TRACE_ROWS, "Tabel 13. Matriks traceability", 8)

    doc.add_heading("12. Risiko dan Rekomendasi", level=1)
    add_doc_table(doc, ["No", "Risiko/Batasan", "Rekomendasi"], [
        ["1", "Backend transaksi belum menolak stok negatif bila API dipanggil langsung tanpa UI POS.", "Tambahkan validasi stok cukup di backend sebelum insert transaksi dan deduct inventory."],
        ["2", "Cash received, customer name, dan payment note pada modal POS dipakai untuk struk UI, tetapi backend saat ini menyimpan amount_paid setara total dan change_amount 0.", "Tambahkan field amount_paid, change_amount, customer_name, dan payment_note pada payload transaksi bila data tersebut harus muncul di laporan."],
        ["3", "Update recipe memakai ON DUPLICATE KEY; skema product_ingredients saat ini belum terlihat memiliki unique key gabungan product_id dan ingredient_id.", "Tambahkan unique index product_id, ingredient_id agar upsert backend konsisten."],
        ["4", "Aplikasi bergantung pada XAMPP lokal.", "Siapkan checklist start Apache/MySQL dan backup database sebelum demo."],
    ])

    doc.add_heading("13. Penutup", level=1)
    doc.add_paragraph("SKPL ini mencerminkan kondisi aplikasi Locales Pro saat ini: POS tunai/manual dengan aktor admin, kasir, dan pembeli, registrasi kasir, manajemen cabang, produk, resep, bahan baku, stok, laporan, void, dan audit mutasi stok. Dokumen ini dapat dipakai sebagai dasar pengujian dan pengembangan lanjutan.")
    doc.save(REPORT_DOCX)


def main():
    ASSET_DIR.mkdir(exist_ok=True)
    generate_use_case_diagram()
    generate_architecture_diagram()
    generate_sequence_diagram()
    generate_data_model_diagram()
    build_markdown()
    build_docx()
    print(f"Generated {REPORT_MD.name}")
    print(f"Generated {REPORT_DOCX.name}")
    print(f"Generated diagrams in {ASSET_DIR.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
